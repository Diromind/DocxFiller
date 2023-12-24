from docx import Document
from time import sleep
import datetime


if __name__ == '__main__':
    doc_template = Document("data\\template.docx")
    par_template = doc_template.paragraphs[3]

    doc_output = Document("output.docx")
    par_output = doc_output.paragraphs[3]

    log = open("data\\debug_log.txt", "a", encoding="utf-8")

    print("Запись от", datetime.datetime.now(), file=log)
    print("Разложение по runs 4го параграфа документа template.docx")
    for run in par_template.runs:
        print(run.text, end="|")
    print("\n---------------------------------")
    print("Разложение по runs 4го параграфа документа output.docx")
    for run_out in par_output.runs:
        print(run_out.text, end="|")
    print("\n---------------------------------")

    print("Разложение по runs 4го параграфа документа template.docx", file=log)
    for run in par_template.runs:
        print(run.text, end="|", file=log)
    print("\n---------------------------------", file=log)
    print("Разложение по runs 4го параграфа документа output.docx", file=log)
    for run_out in par_output.runs:
        print(run_out.text, end="|", file=log)
    print("\n---------------------------------", file=log)

    log.close()

    print("Спасибо!")
    sleep(2)
