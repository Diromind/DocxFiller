import sys
import datetime
import locale
import typing as tp
from copy import deepcopy


import json

from PyQt6.QtCore import QRect, Qt
from PyQt6.QtGui import QIcon, QGuiApplication
from PyQt6.QtWidgets import QApplication, QWidget, QVBoxLayout, QHBoxLayout, QLabel, QLineEdit, QPushButton

from docx import Document
locale.setlocale(locale.LC_ALL, ('ru_RU', 'UTF-8'))


class TypeOfDataField:
    def __init__(self,
                 placeholder,
                 display_name: str = None,
                 default_value: str = None,
                 group: str = None) -> None:

        self.placeholder = placeholder
        self.display_name = display_name or placeholder
        self.default_value = default_value

        self.widget = None
        self.value = None

        self.group = group

    def __str__(self):
        if self.value is None:
            return self.display_name + " : None, default set: " + self.default_value
        else:
            return self.display_name + " : " + str(self.value)


class Settings:
    def __init__(self) -> None:
        self.screen = QGuiApplication.primaryScreen().geometry()
        self.geometry = self.set_default_geometry()
        self.spacing = 4

    def set_default_geometry(self) -> QRect:
        return QRect(self.screen.width() // 4, self.screen.height() // 4,
                     self.screen.width() // 2, self.screen.height() // 2)


class Controller:
    def __init__(self,
                 path_to_data_folder: str = ""):

        self.data_folder_path = path_to_data_folder + "\\"
        self.app = QApplication([])
        self.filler = DocumentFiller(self.data_folder_path + "template.docx")

        self._apply_settings_()
        self._set_current_date_()

        self.fields_data = [TypeOfDataField("__НомерДоговора__", "Номер договора", "46"),
                            TypeOfDataField("__Дата__", "Дата", self.date_short),
                            TypeOfDataField("__ДатаБуквами__", "Дата прописная", self.date_long),
                            TypeOfDataField("__СтоимостьРаботы__", "Стоимость работы", "47000 (сорок семь тысяч) рублей"),
                            TypeOfDataField("__КлассыМКТУ__", "Классы МКТУ, через запятую", "25, 28"),
                            TypeOfDataField("__ПолноеИмя__", "ФИО Клиента", "Фамилия Имя Отчество ", "personal_data"),
                            TypeOfDataField("__ЭлектроннаяПочта__", "Email", "example@yandex.ru", "personal_data"),
                            TypeOfDataField("__Адрес__", "Адрес Клиента", "123456, Московская область, "
                                                                          "г. Долгопрудный, "
                                                                          "Московское шоссе, д. 21 к. 6", "personal_data"),
                            TypeOfDataField("__АдресДляПереписки__", "Адрес для переписки с Роспатентом", "123456, "
                                                                        "Московская область, г. Долгопрудный, "
                                                                        "Московское шоссе, д. 21 к. 6", "personal_data"),
                            TypeOfDataField("__НомерПаспорта__", "Номер Паспорта", "1234 567890", "personal_data"),
                            TypeOfDataField("__ПаспортВыданКем__", "Орган, выдавший паспорт", "выдан МВД по городу "
                                                                                              "Москва",
                                            "personal_data"),
                            TypeOfDataField("__ПаспортВыданДата__", "Дата выдачи паспорта", "12.34.5678",
                                            "personal_data"),
                            TypeOfDataField("__ПаспортКодПодразделения__", "Код подразделения", "123-456",
                                            "personal_data"),
                            TypeOfDataField("__ПаспортАдрес__", "Адрес по паспорту", "123456, Московская область, "
                                                                          "г. Долгопрудный, "
                                                                          "Московское шоссе, д. 21 к. 6", "personal_data"),
                            TypeOfDataField("__ИНН__", "Номер ИНН", "123456789012", "personal_data"),
                            TypeOfDataField("__СНИЛС__", "СНИЛС", "123-456-789 00", "personal_data")]
        self.output_file = TypeOfDataField(None, "Имя файла, куда будет сохранен результат", "output.docx")
        self.output_file.value = "output.docx"

    def _apply_settings_(self) -> None:
        self.settings = Settings()
        self.app.setWindowIcon(QIcon(self.data_folder_path + "icon.png"))

    def run(self) -> None:
        window = RenderApp(self, self.settings, self.fields_data, self.output_file)
        window.setWindowTitle('Little Helper')
        window.setGeometry(self.settings.geometry)
        window.show()
        sys.exit(self.app.exec())

    def save_inputs_to_file(self) -> None:
        self.filler.update_data(self.fields_data)
        self.filler.fill_data()
        self.filler.save(self.output_file.value)

    def _set_current_date_(self) -> None:
        self.months_with_correct_spelling = {1: "января", 2: "февраля", 3: "марта", 4: "апреля",
                                             5: "мая", 6: "июня", 7: "июля", 8: "августа",
                                             9: "сентября", 10: "октября", 11: "ноября", 12: "декабря"}
        self.date_short = datetime.datetime.now().strftime("%d.%m.%Y") + " г."
        date_long_template = datetime.datetime.now().strftime("%d __month__ %Y") + " г."

        correct_spelled_month = self.months_with_correct_spelling[datetime.datetime.now().month]
        self.date_long = date_long_template.replace("__month__", correct_spelled_month)


class DocumentFiller:
    def __init__(self,
                 template_name: str = "template.docx",
                 fields_data=None):
        self.template = Document(template_name)
        self.output_document = deepcopy(self.template)
        self.field_data = deepcopy(fields_data)

    def update_data(self, new_data):
        self.field_data = new_data

        self.output_document = deepcopy(self.template)

    def fill_data(self):
        for field in self.field_data:
            if field.value is not None:
                self._replace_placeholder_(field.placeholder, field.value)
            elif field.default_value is not None:
                self._replace_placeholder_(field.placeholder, field.default_value)

    def _replace_placeholder_(self, placeholder, value):
        for paragraph in self.output_document.paragraphs:
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)
        for table in self.output_document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, value)

    def save(self, output_name="output.docx") -> None:
        try:
            self.output_document.save(output_name)
            with open("data\\log.txt", "a", encoding='utf-8') as f:
                print("Файл успешно сохранен", file=f)
                print("Значения полей:\n", *controller.fields_data, sep="  |||  ", file=f, end="\n\n")
        except:
            with open("data\\log.txt", "a", encoding='utf-8') as f:
                print("Не удалось сохранить файл. Возможно, файл с таким названием открыт: "
                      "закройте его и повторите попытку", file=f)


class RenderApp(QWidget):
    def __init__(self,
                 controller_instance: Controller,
                 settings_set: Settings,
                 fields: list,
                 output_file):
        super().__init__()
        self.controller = controller_instance
        self.setStyleSheet("QLineEdit { color: #336699; }")
        self.settings = settings_set

        self.fields = fields
        self.output_file = output_file

        self.initUI()

    def initUI(self):
        main_layout = QHBoxLayout()

        layout_personal = QVBoxLayout()
        layout_document_wise = QVBoxLayout()
        layout_personal.setSpacing(self.settings.spacing)
        layout_document_wise.setSpacing(self.settings.spacing)
        layout_personal.setAlignment(Qt.AlignmentFlag.AlignTop)
        layout_document_wise.setAlignment(Qt.AlignmentFlag.AlignTop)

        for field in self.fields:
            if field.group == "personal_data":
                self.create_input_field(field, layout_personal)
            else:
                self.create_input_field(field, layout_document_wise)

        layout_document_wise.addSpacing(45)
        self.create_input_field(self.output_file, layout_document_wise)

        main_layout.addLayout(layout_document_wise)
        main_layout.addLayout(layout_personal)

        self.setLayout(main_layout)

        save_button = QPushButton("Сохранить")
        save_button.setFixedSize(100, 30)
        save_button.clicked.connect(self.controller.save_inputs_to_file)
        layout_document_wise.addWidget(save_button, alignment=Qt.AlignmentFlag.AlignHCenter)

    def create_input_field(self, field, layout):
        label = QLabel(f'{field.display_name}')
        text_input_widget = QLineEdit()
        text_input_widget.setPlaceholderText(field.default_value)

        text_input_widget.editingFinished.connect(lambda field_finished=field: self.save_field_data(field_finished))
        field.widget = text_input_widget

        layout.addWidget(label)
        layout.addWidget(field.widget)

    @staticmethod
    def save_field_data(field):
        field.value = field.widget.text()
        with open("data\\log.txt", "a", encoding='utf-8') as f:
            print(f"Сохранено значение поля {field.display_name}: {field.widget.text()}", file=f)

    def moveEvent(self, event):
        self.settings.geometry.x = event.pos().x
        self.settings.geometry.y = event.pos().y

    def resizeEvent(self, event):
        self.settings.geometry.width = event.size().width()
        self.settings.geometry.height = event.size().height()

    def closeEvent(self, event):
        self.controller.save_inputs_to_file()
        # Save window size to settings when the window is closed
        settings = {"window_size": {"width": self.width(), "height": self.height()}}
        with open("data\\log.txt", "a", encoding='utf-8') as f:
            print("Session end, time is", datetime.datetime.now(), file=f)



if __name__ == '__main__':
    with open("data\\log.txt", "a", encoding='utf-8') as log_f:
        print("----------------------------------------------------------------\n"
              "Start session, time is", datetime.datetime.now(), file=log_f)
    controller = Controller("data")
    controller.run()





